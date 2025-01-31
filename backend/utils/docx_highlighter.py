from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import logging
import json
from pathlib import Path

logger = logging.getLogger(__name__)


def get_highlight_color(color_name: str) -> WD_COLOR_INDEX:
    """Get highlight color for a given color name"""
    colors = {
        "green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "yellow": WD_COLOR_INDEX.YELLOW,
        "pink": WD_COLOR_INDEX.PINK,
        "blue": WD_COLOR_INDEX.TURQUOISE,
    }
    return colors.get(color_name, WD_COLOR_INDEX.YELLOW)  # Default to yellow


def add_highlights(highlight_data: list, filename: str) -> str:
    """
    Add color-coded highlights to a DOCX file.

    Args:
        highlight_data (list): List of dictionaries containing content and highlight information
        filename (str): Path to the DOCX file

    Returns:
        str: Path to the highlighted DOCX file
    """
    logger.info("Adding highlights to %s", filename)

    try:
        # Open document
        doc = Document(filename)

        # Process each highlight
        for highlight in highlight_data:
            content = highlight["content"]
            color = get_highlight_color(highlight["highlight_type"]["color"])

            # Search through all paragraphs
            for paragraph in doc.paragraphs:
                if content in paragraph.text:
                    # Split text and apply highlighting
                    start_idx = paragraph.text.index(content)
                    end_idx = start_idx + len(content)

                    # Clear existing runs
                    for run in paragraph.runs:
                        run.clear()

                    # Add text before highlight
                    if start_idx > 0:
                        run = paragraph.add_run(paragraph.text[:start_idx])

                    # Add highlighted text
                    run = paragraph.add_run(content)
                    run.font.highlight_color = color

                    # Add comment if supported by python-docx
                    # Currently python-docx doesn't support comments directly

                    # Add text after highlight
                    if end_idx < len(paragraph.text):
                        run = paragraph.add_run(paragraph.text[end_idx:])

        # Save highlighted document
        output_path = str(Path("files") / f"highlighted_{Path(filename).name}")
        doc.save(output_path)
        logger.info("Saved highlighted file to %s", output_path)
        return output_path

    except Exception as e:
        logger.error("Error highlighting DOCX: %s", str(e))
        return None
